# core/ui/evidence_viewer.py
from __future__ import annotations

import io
import json
import time
import zipfile
from datetime import datetime, timezone
from typing import Any, Dict, Optional

import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from core.classify.detect import detect_cui
from core.evidence.query import build_query_service
from core.utils.hash_utils import sha256_bytes_hex


# ----------------------------
# HELPERS
# ----------------------------
def _ms_to_iso(ms: Optional[int]) -> str:
    if not ms:
        return ""
    try:
        return datetime.fromtimestamp(ms / 1000.0, tz=timezone.utc).isoformat()
    except Exception:
        return ""


def _safe_json_loads(s: Optional[str]):
    if not s:
        return None
    if isinstance(s, dict):
        return s
    try:
        return json.loads(s)
    except Exception:
        return s


def get_detection_matches(detection: Optional[dict]) -> list[str]:
    """Support both detector formats without changing detect_cui()."""
    if not detection:
        return []

    matches = detection.get("matches")
    if matches:
        if isinstance(matches, dict):
            flat = []
            for value in matches.values():
                if isinstance(value, list):
                    flat.extend(value)
                elif isinstance(value, str):
                    flat.append(value)
            return [str(m) for m in flat if m]
        if isinstance(matches, list):
            return [str(m) for m in matches if m]
        if isinstance(matches, str):
            return [matches]

    return [
        str(hit.get("match"))
        for hit in detection.get("rule_hits", []) or []
        if isinstance(hit, dict) and hit.get("match")
    ]


def _get_child_attachments(storage, parent_evidence_id: str):
    with storage.ledger._connect() as con:
        rows = con.execute("""
            SELECT *
            FROM evidence_records
            WHERE json_extract(metadata_json, '$.parent_email_evidence_id') = ?
        """, (parent_evidence_id,)).fetchall()

    return [dict(r) for r in rows]


def _highlight_cui(text: str, detection: dict):
    if not text:
        return text

    matches = get_detection_matches(detection)
    matches = sorted(set(matches), key=len, reverse=True)

    for m in matches:
        if not m or len(m) < 3:
            continue
        try:
            text = text.replace(
                m,
                f"<mark style='background-color:#ff4d4f;color:white;padding:2px 4px;border-radius:3px'>{m}</mark>"
            )
        except Exception:
            pass

    return text


def _build_case_report(storage, case_id):
    ledger = storage.ledger

    if hasattr(ledger, "get_case_details"):
        data = ledger.get_case_details(case_id) or {}

        timeline = data.get("timeline") or []

        # =======================================
        # 🔥 FIX: BACKFILL TIMELINE IF EMPTY
        # =======================================
        if not timeline:
            try:
                with ledger._connect() as con:
                    custody = con.execute("""
                        SELECT ce.*
                        FROM custody_events ce
                        JOIN case_evidence cex ON ce.evidence_id = cex.evidence_id
                        WHERE cex.case_id = ?
                        ORDER BY ce.timestamp_ms ASC
                    """, (case_id,)).fetchall()

                timeline = [dict(c) for c in custody]
            except Exception as e:
                print("⚠️ Timeline backfill failed:", e)
        # =======================================

        return {
            "case": data.get("case") or {},
            "alerts": data.get("alerts") or [],
            "evidence": data.get("evidence") or [],
            "timeline": timeline,
        }

    case = next(
        (c for c in ledger.list_cases() if (c.get("id") or c.get("case_id")) == case_id),
        None,
    )
    evidence = ledger.list_case_evidence(case_id) if hasattr(ledger, "list_case_evidence") else []
    timeline = ledger.build_case_timeline(case_id) if hasattr(ledger, "build_case_timeline") else []

    alerts = []
    if hasattr(ledger, "list_alerts"):
        all_alerts = ledger.list_alerts(limit=500)
        alerts = [a for a in all_alerts if a.get("case_id") == case_id]

    return {
        "case": case,
        "alerts": alerts,
        "evidence": evidence,
        "timeline": timeline,
    }


def _get_case_id_for_evidence(storage, evidence_id):
    with storage.ledger._connect() as con:
        row = con.execute("""
            SELECT case_id
            FROM case_evidence_map
            WHERE evidence_id = ?
        """, (evidence_id,)).fetchone()

    return row[0] if row else None


def ensure_custody(storage, evidence_id):
    with storage.ledger._connect() as con:
        row = con.execute("""
            SELECT 1 FROM custody_events
            WHERE evidence_id = ?
            LIMIT 1
        """, (evidence_id,)).fetchone()

        if not row:
            con.execute("""
                INSERT INTO custody_events (
                    run_id,
                    evidence_id,
                    event_type,
                    actor,
                    timestamp_ms,
                    details_json
                )
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                94,
                evidence_id,
                "INGESTED",
                "auto_heal",
                int(time.time() * 1000),
                json.dumps({"auto_healed": True})
            ))
            con.commit()


def _generate_pdf_report(report, output_path="forensic_report.pdf"):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(output_path, pagesize=letter)

    elements = []
    elements.append(Paragraph("Forensic Evidence Report", styles["Title"]))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(
        f"Generated: {datetime.utcnow().isoformat()} UTC",
        styles["Normal"]
    ))
    elements.append(Spacer(1, 20))

    report_bytes = json.dumps(report, sort_keys=True, default=str).encode()
    report_hash = sha256_bytes_hex(report_bytes)

    elements.append(Paragraph("Integrity Hash (SHA256):", styles["Heading3"]))
    elements.append(Paragraph(report_hash, styles["Code"]))
    elements.append(Spacer(1, 20))

    table_data = [["Evidence ID", "Type", "Size", "Flags"]]

    # Build lookup: evidence_id → alert detections
    alert_map = {}

    for a in report.get("alerts", []):
        eid = a.get("evidence_id")

        try:
            notes = json.loads(a.get("notes") or "{}")
        except Exception:
            notes = {}

        categories = notes.get("categories") or notes.get("matches") or []

        if eid and categories:
            alert_map.setdefault(eid, set()).update(categories)

    # Build table
    for e in report.get("evidence", []) or []:
        eid = e.get("evidence_id")

        detections = list(alert_map.get(eid, []))

        # 🔥 fallback: handle string mismatch / whitespace / casing
        if not detections:
            for k, v in alert_map.items():
                if str(k).strip().lower() == str(eid).strip().lower():
                    detections = list(v)
                    break

        table_data.append([
            str(eid or ""),
            str(e.get("content_type") or ""),
            str(e.get("size_bytes") or ""),
            ", ".join(sorted(detections)) if detections else "NONE",
        ])

    table = Table(table_data)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 20))

    elements.append(Paragraph("Chain of Custody", styles["Heading2"]))
    for c in report.get("timeline", []) or []:
        timestamp = c.get("timestamp") or c.get("timestamp_ms") or c.get("created_at_ms") or ""
        event_type = c.get("event_type") or c.get("type") or ""
        actor = c.get("actor") or ""
        elements.append(Paragraph(f"{timestamp} | {event_type} | {actor}", styles["Normal"]))

    doc.build(elements)
    return output_path


def export_evidence_bundle(storage, case_id):
    vault = storage.vault
    ledger = storage.ledger

    with ledger._connect() as con:
        evidence = con.execute("""
            SELECT e.*
            FROM evidence_records e
            JOIN case_evidence ce ON e.evidence_id = ce.evidence_id
            WHERE ce.case_id = ?
        """, (case_id,)).fetchall()

        custody = con.execute("""
            SELECT ce.*
            FROM custody_events ce
            JOIN case_evidence cex ON ce.evidence_id = cex.evidence_id
            WHERE cex.case_id = ?
            ORDER BY ce.timestamp_ms ASC
        """, (case_id,)).fetchall()

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as z:
        for e in evidence:
            e = dict(e)
            eid = e["evidence_id"]
            try:
                data = vault.open_bytes(evidence_id=eid)
                z.writestr(f"evidence/{eid}.bin", data)
            except Exception:
                pass

        z.writestr("metadata/evidence.json", json.dumps([dict(e) for e in evidence], indent=2, default=str))
        z.writestr("metadata/custody.json", json.dumps([dict(c) for c in custody], indent=2, default=str))

    zip_buffer.seek(0)
    return zip_buffer


def _render_export_section(storage: Any, evidence_id: str) -> None:
    st.divider()
    st.subheader("📤 Export Forensic Report")

    case_id = _get_case_id_for_evidence(storage, evidence_id)

    if not case_id:
        st.warning("No case associated with this evidence.")
        return

    report = _build_case_report(storage, case_id)
    json_bytes = json.dumps(report, indent=2, default=str).encode("utf-8")

    st.download_button(
        label="⬇️ Download JSON Report",
        data=json_bytes,
        file_name=f"case_{case_id}.json",
        mime="application/json",
    )

    try:
        pdf_path = _generate_pdf_report(report)
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        st.download_button(
            label="⬇️ Download PDF Report",
            data=pdf_bytes,
            file_name=f"case_{case_id}_report.pdf",
            mime="application/pdf",
        )
    except Exception as e:
        st.warning(f"PDF generation failed: {e}")

    st.markdown("### 📦 Export Evidence Bundle")

    try:
        zip_data = export_evidence_bundle(storage, case_id)
        st.download_button(
            label="⬇️ Download Full Evidence Bundle",
            data=zip_data,
            file_name=f"case_{case_id}_bundle.zip",
            mime="application/zip"
        )
    except Exception as e:
        st.error(f"Bundle export failed: {e}")


# ----------------------------
# MAIN VIEW
# ----------------------------
def render_evidence_viewer(storage: Any) -> None:

    query = build_query_service(storage)
    ledger = getattr(storage, "ledger", None)
    vault = getattr(storage, "vault", None)

    st.markdown("## 🧾 Evidence Viewer")

    if ledger is None:
        st.error("Storage ledger not available.")
        return

    if vault is None:
        st.error("Storage vault not available.")
        return

    immutable = bool(getattr(ledger, "immutable", False))
    st.success("🛡 Immutable Ledger Mode ENABLED") if immutable else st.warning("⚠ Mutable Ledger Mode")

    # =========================================================
    # 🔗 DIRECT MODE (FROM ALERT CLICK)
    # =========================================================
    preselected_evidence_id = st.session_state.get("selected_evidence_id")

    if preselected_evidence_id:

        st.info(f"🔗 Linked Evidence Loaded: {preselected_evidence_id}")

        evidence_id = str(preselected_evidence_id)

        # ----------------------------
        # LOAD RECORD (STRICT MATCH)
        # ----------------------------
        record = None

        try:
            if hasattr(ledger, "get_evidence_record"):
                record = ledger.get_evidence_record(evidence_id)

            if not record:
                with storage.ledger._connect() as con:
                    row = con.execute("""
                        SELECT evidence_id, sha256, size_bytes, content_type, suggested_name, metadata_json
                        FROM evidence_records
                        WHERE evidence_id = ?
                        LIMIT 1
                    """, (evidence_id,)).fetchone()

                    if row:
                        record = dict(row)

        except Exception as e:
            st.error(f"Error loading record: {e}")
            return

        if not record:
            # -------------------------------------------------
            # FALLBACK: Evidence exists in vault but not ledger
            # -------------------------------------------------
            try:
                data = vault.open_bytes(evidence_id=evidence_id)

                if data:
                    record = {
                        "evidence_id": evidence_id,
                        "sha256": sha256_bytes_hex(data),
                        "size_bytes": len(data),
                        "content_type": "application/octet-stream",
                        "source": "vault_fallback",
                    }
                else:
                    st.error("Linked evidence not found in ledger or vault.")
                    return

            except Exception as e:
                st.error(f"Linked evidence not found in ledger or vault: {e}")
                return

        # ----------------------------
        # DISPLAY RECORD
        # ----------------------------
        st.markdown("### 📄 Evidence Record")
        st.json(record)
        # ---------------------------------------
        # 📊 INGEST COUNT (SAFE ADD)
        # ---------------------------------------
        ingest_count = 0

        try:
            with storage.ledger._connect() as con:
                row = con.execute("""
                    SELECT COUNT(*)
                    FROM custody_events
                    WHERE evidence_id = ?
                    AND event_type = 'INGESTED'
                """, (evidence_id,)).fetchone()

                if row:
                    ingest_count = row[0]

        except Exception as e:
            st.warning(f"Failed to load ingest count: {e}")

        if ingest_count > 0:
            st.info(f"📊 Seen {ingest_count} time{'s' if ingest_count != 1 else ''}")
        # ----------------------------
        # CONTENT PREVIEW
        # ----------------------------
        st.markdown("### 🔎 Content Preview")

        try:
            data = vault.open_bytes(evidence_id=evidence_id)

            content_type = (record.get("content_type") or "").lower()
            name = (record.get("suggested_name") or "").lower()

            if not name:
                name = f"{evidence_id}.bin"

            # DOCX detection
            if data[:2] == b"PK":
                try:
                    with zipfile.ZipFile(io.BytesIO(data)) as z:
                        if any("word/" in f for f in z.namelist()):
                            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            name = f"{evidence_id}.docx"
                except Exception:
                    pass

            # PDF detection
            elif data[:4] == b"%PDF":
                content_type = "application/pdf"
                name = f"{evidence_id}.pdf"

            # ---------------------------------------
            # 🧠 UNIFIED TEXT EXTRACTION
            # ---------------------------------------
            text = ""

            try:
                if name.endswith(".docx") or "wordprocessingml" in content_type:
                    doc = Document(io.BytesIO(data))
                    text = "\n".join(p.text for p in doc.paragraphs)

                elif name.endswith(".pdf") or "pdf" in content_type:
                    reader = PdfReader(io.BytesIO(data))
                    text = "\n".join([p.extract_text() or "" for p in reader.pages])

                elif name.endswith(".txt") or "text" in content_type:
                    text = data.decode("utf-8", errors="ignore")

                else:
                    # Safe fallback for text-like .bin evidence.
                    decoded = data.decode("utf-8", errors="ignore")
                    text = decoded if decoded.strip() else ""

            except Exception as e:
                st.warning(f"Text extraction failed: {e}")
                text = ""

            # ---------------------------------------
            # 🔥 RUN DETECTION (ALWAYS AFTER EXTRACTION)
            # ---------------------------------------
            if text.strip():
                try:
                    detection = detect_cui(text)
                except Exception as e:
                    st.warning(f"Detection failed: {e}")
                    detection = {"matches": []}
            else:
                detection = {"matches": []}

            # ---------------------------------------
            # 🔍 DISPLAY CONTENT
            # ---------------------------------------
            if text.strip():
                highlighted = _highlight_cui(text[:5000], detection)
                st.markdown(highlighted, unsafe_allow_html=True)
            else:
                st.info("No preview available for this file type.")

            # ---------------------------------------
            # 🧠 DETECTION SUMMARY
            # ---------------------------------------
            st.markdown("### 🧠 Detected CUI")

            matches = get_detection_matches(detection)

            if matches:
                st.error(f"⚠️ DETECTED: {', '.join(sorted(set(matches)))}")
            else:
                st.success("No CUI detected")

        except Exception as e:
            st.error(f"Preview error: {e}")

        # ----------------------------
        # 🧾 CHAIN OF CUSTODY
        # ----------------------------
        st.markdown("### 📜 Chain-of-Custody Timeline")

        eid = (record.get("evidence_id") or "").strip()
        st.write("DEBUG EVIDENCE ID USED:", eid)

        with storage.ledger._connect() as con:
            rows = con.execute("""
                SELECT event_type, actor, timestamp_ms, details_json
                FROM custody_events
                WHERE TRIM(evidence_id) = ?
                ORDER BY timestamp_ms ASC
            """, (eid,)).fetchall()

        if rows:
            for r in rows:
                ts = r[2]
                event = r[0]
                actor = r[1]
                st.write(f"🕒 {ts} | {event} | {actor}")
        else:
            st.warning("No custody events found.")

        # ----------------------------
        # VERIFY
        # ----------------------------
        st.divider()
        st.markdown("### 🔐 Verify")

        try:
            data = vault.open_bytes(evidence_id=evidence_id)
            actual = sha256_bytes_hex(data)
            expected = record.get("sha256")

            if actual == expected:
                st.success("✅ VERIFIED")
            else:
                st.error("❌ TAMPERED")

        except Exception as e:
            st.error(f"Verification error: {e}")

        _render_export_section(storage, evidence_id)

        # ----------------------------
        # CLEAR BUTTON
        # ----------------------------
        if st.button("↩ Clear Linked Evidence", use_container_width=True):
            st.session_state.pop("selected_evidence_id", None)
            st.session_state.pop("alert_notes", None)
            st.session_state["page"] = "Scan"
            st.rerun()

        return

    # =========================================================
    # STANDARD MODE (NO DIRECT LINK)
    # =========================================================
    runs = query.list_recent_runs(limit=50)

    if not runs:
        st.info("No runs found yet.")
        return

    labels = [
        f"{r.get('started_at_utc', '')} | {r.get('mailbox', '')}"
        for r in runs
    ]

    idx = st.selectbox("Select Run", options=list(range(len(labels))), format_func=lambda i: labels[i])
    run_id = runs[idx]["run_id"]

    st.subheader("Run Summary")
    st.json(query.load_run_summary(run_id) or {})

    st.subheader("Manifest")
    st.json(query.load_manifest(run_id) or {})

    evidence = query.list_evidence_for_run(run_id)

    if not evidence:
        st.info("No evidence found.")
        return

    ev_labels = [
        f"{e.get('suggested_name','(unnamed)')} | {str(e.get('evidence_id',''))[:12]}"
        for e in evidence
    ]

    ev_idx = st.selectbox("Select Evidence", options=list(range(len(ev_labels))), format_func=lambda i: ev_labels[i])
    record: Dict[str, Any] = evidence[ev_idx]
    evidence_id = str(record.get("evidence_id"))

    # ---------------------------------------
    # 🔗 LOAD ATTACHMENTS
    # ---------------------------------------
    attachments = _get_child_attachments(storage, evidence_id)
    print("DEBUG ATTACHMENTS FOUND:", len(attachments))

    # ---------------------------------------
    # 🔐 VERIFY FUNCTION
    # ---------------------------------------
    def _verify(eid, rec):
        try:
            data = vault.open_bytes(evidence_id=eid)
            actual = sha256_bytes_hex(data)
            expected = rec.get("sha256")
            return actual == expected
        except Exception:
            return False

    # ---------------------------------------
    # VERIFY SUMMARY
    # ---------------------------------------
    total = 1 + len(attachments)
    verified = sum(
        [_verify(evidence_id, record)] +
        [_verify(a["evidence_id"], a) for a in attachments]
    )

    st.info(f"Integrity Score: {verified}/{total} objects verified")

    # VERIFY
    st.markdown("### 🔐 Verify")

    # ---------------------------------------
    # VERIFY EMAIL
    # ---------------------------------------
    email_ok = _verify(evidence_id, record)

    if email_ok:
        st.success("✅ Email VERIFIED")
    else:
        st.error("❌ Email TAMPERED")

    # ---------------------------------------
    # VERIFY + ANALYZE ATTACHMENTS
    # ---------------------------------------
    if attachments:
        st.markdown("#### 📎 Attachment Verification")

        verified_count = 0

        for att in attachments:
            name = att.get("suggested_name") or "attachment"
            eid = att.get("evidence_id")

            # -----------------------------
            # 🔐 VERIFY
            # -----------------------------
            try:
                ok = _verify(eid, att)

                if ok:
                    verified_count += 1
                    st.success(f"✅ {name} VERIFIED")
                else:
                    st.error(f"❌ {name} TAMPERED")

            except Exception as e:
                st.warning(f"Verification failed for {name}: {e}")
                continue

            # -----------------------------
            # 🧠 LOAD ATTACHMENT DATA
            # -----------------------------
            try:
                att_data = vault.open_bytes(evidence_id=eid)

                att_name = (att.get("suggested_name") or "").lower()
                att_type = (att.get("content_type") or "").lower()

                # Detect DOCX (zip-based)
                if att_data[:2] == b"PK":
                    try:
                        with zipfile.ZipFile(io.BytesIO(att_data)) as z:
                            if any("word/" in f for f in z.namelist()):
                                att_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                if not att_name:
                                    att_name = "detected.docx"
                    except Exception:
                        pass

                # Detect PDF
                elif att_data[:4] == b"%PDF":
                    att_type = "application/pdf"
                    if not att_name:
                        att_name = "detected.pdf"

            except Exception as e:
                st.warning(f"Failed to load {att.get('suggested_name')}: {e}")
                continue

            # -----------------------------
            # 🧠 EXTRACT TEXT (ATTACHMENTS)
            # -----------------------------
            att_text = ""

            try:
                if att_name.endswith(".docx") or "wordprocessingml" in att_type:
                    doc = Document(io.BytesIO(att_data))
                    att_text = "\n".join(p.text for p in doc.paragraphs)

                elif att_name.endswith(".pdf") or "pdf" in att_type:
                    reader = PdfReader(io.BytesIO(att_data))
                    att_text = "\n".join([p.extract_text() or "" for p in reader.pages])

                elif att_name.endswith(".txt") or "text" in att_type:
                    att_text = att_data.decode("utf-8", errors="ignore")

                else:
                    decoded = att_data.decode("utf-8", errors="ignore")
                    if decoded.strip():
                        att_text = decoded
                    else:
                        st.info(f"Skipping unsupported attachment type: {att_name}")
                        continue

            except Exception as e:
                st.warning(f"Text extraction failed for {att_name}: {e}")
                continue

            # -----------------------------
            # 🔍 SHOW CLEAN PREVIEW
            # -----------------------------
            if att_text.strip():
                with st.expander(f"🔍 Preview: {name}"):
                    st.markdown(att_text[:3000])

            # -----------------------------
            # 🔥 ATTACHMENT DETECTION
            # -----------------------------
            attachment_cui_detected = False

            if att_text.strip():
                try:
                    att_detection = detect_cui(att_text)
                    matches = get_detection_matches(att_detection)

                    if matches:
                        attachment_cui_detected = True
                        st.error(f"🚨 {att_name} DETECTED: {', '.join(sorted(set(matches)))}")

                        highlighted = _highlight_cui(att_text[:5000], att_detection)
                        st.markdown(highlighted, unsafe_allow_html=True)

                    else:
                        st.success(f"🟢 {att_name} — No CUI detected")

                except Exception as e:
                    st.warning(f"Detection failed for {att_name}: {e}")

            # ---------------------------------------
            # 🔗 PROPAGATE TO PARENT EVIDENCE
            # ---------------------------------------
            if attachment_cui_detected:
                st.error("🚨 CUI DETECTED IN ATTACHMENTS (propagated to parent evidence)")
                storage.ledger.create_alert(
                    evidence_id=evidence_id,
                    severity="CRITICAL",
                    message="CUI detected in attachment(s)"
                )

    _render_export_section(storage, evidence_id)
