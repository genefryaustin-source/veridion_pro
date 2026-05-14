from io import BytesIO
from typing import Any, Dict

from docx import Document

from core.extractors.base import BaseExtractor
from core.extractors.models import ExtractedContent


class DocxExtractor(BaseExtractor):
    supported_extensions = [".docx"]
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def extract(
        self,
        data: bytes,
        filename: str,
        content_type: str = "",
        metadata: Dict[str, Any] | None = None,
    ) -> ExtractedContent:

        warnings = []
        text_parts = []

        try:
            doc = Document(BytesIO(data))

            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    text_parts.append(p.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    values = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text and cell.text.strip()
                    ]
                    if values:
                        text_parts.append(" | ".join(values))

        except Exception as e:
            warnings.append(str(e))

        text = "\n".join(text_parts).strip()

        return ExtractedContent(
            text=text,
            filename=filename,
            content_type=content_type or self.supported_mime_types[0],
            extension=".docx",
            extraction_method="python_docx",
            confidence="HIGH" if text else "LOW",
            metadata=metadata or {},
            warnings=warnings,
        )